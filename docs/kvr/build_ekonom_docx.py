from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

md_path = r"c:\Users\anton\Desktop\app\docs\kvr\06-razdil-4.md"
out_path = r"c:\Users\anton\Downloads\економічний розділ KeyGames.docx"

with open(md_path, "r", encoding="utf-8") as f:
    text = f.read()
text = text.replace("Здod", "Здод")
with open(md_path, "w", encoding="utf-8") as f:
    f.write(text)

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(3.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(1.0)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)

lines = text.split("\n")
i = 0
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith("# "):
        p = doc.add_paragraph(line[2:])
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith("## "):
        p = doc.add_paragraph(line[3:])
        p.runs[0].bold = True
    elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(row)
            i += 1
        if len(rows) >= 2 and set("".join(rows[1])) <= set("|-: "):
            rows = [rows[0]] + rows[2:]
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                table.rows[r].cells[c].text = val
        continue
    elif line.startswith("- "):
        doc.add_paragraph(line[2:], style="List Bullet")
    elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
        doc.add_paragraph(line.strip("*"))
    else:
        clean = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
        doc.add_paragraph(clean)
    i += 1

doc.save(out_path)
print(f"Saved: {out_path}")
